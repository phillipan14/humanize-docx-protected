"""Post-humanizer cleanup: restore terminology drift introduced by roundtrip translation.

Finds and replaces known drift patterns in the humanized docx so that thesis-critical
terminology is consistent with the original.

Patterns handled:
- Tier 1/2/3 → Layer 1/2/3 (H-S-D framework term)
- "remote operation" → "teleoperation" (only when context matches)
- Capitalized drift ("Key/official" → "Primary/official")
- Contextual robots ↔ robotics (heuristic — only where "robotics industry/ecosystem/market" pattern expected)
"""
import sys, re
from pathlib import Path
from docx import Document

# Direct replacements — safe to apply globally
DIRECT_REPLACEMENTS = [
    # H-S-D layer restoration
    (r'\bTier 1\b', 'Layer 1'),
    (r'\bTier 2\b', 'Layer 2'),
    (r'\bTier 3\b', 'Layer 3'),
    (r'\bFirst layer\b', 'Layer 1'),
    (r'\bSecond layer\b', 'Layer 2'),
    (r'\bThird layer\b', 'Layer 3'),
    (r'\bfirst tier\b', 'Layer 1'),
    (r'\bsecond tier\b', 'Layer 2'),
    (r'\bthird tier\b', 'Layer 3'),
    # Skylarq misspellings from translation drift
    (r'\bSkylark AI\b', 'Skylarq AI'),
    (r'\bSkylark\b', 'Skylarq'),
    # Specific evidence-hierarchy heading drift
    (r'\bKey/official\b', 'Primary/official'),
    (r'Industry Reports and Standards Authority', 'Industry reports and standards bodies'),
    (r'Standards Authority', 'standards bodies'),
    # H-S-D variant spellings
    (r'\bHSD framework\b', 'H-S-D framework'),
    (r'\bH S D framework\b', 'H-S-D framework'),
    # Common translation artifacts
    (r'\bSouth Asia Asia\b', 'Southeast Asia'),
    (r'\bSoutheast Asia Asia\b', 'Southeast Asia'),
    # Hyphenated terms that get split
    (r'\bgo to market\b', 'go-to-market'),
    (r'\bGo to market\b', 'Go-to-market'),
    (r'\bcross border\b', 'cross-border'),
    (r'\bCross border\b', 'Cross-border'),
    (r'\bafter sales\b', 'after-sales'),
    (r'\bAfter sales\b', 'After-sales'),
    # Capitalization fixes
    (r'\bunitedstates\b', 'United States'),
    # Specific academic terms
    (r'\bmachine arm industry\b', 'robotic arm industry'),
    (r'\bmachine arms\b', 'robotic arms'),
    # "teleoperation" variants (very careful — only context-safe)
    (r'\bremote controlOperation\b', 'teleoperation'),
    (r'\bremote operationDataData\b', 'teleoperation data'),
]

# Context-sensitive replacements — use small regex patterns to catch drift
CONTEXT_REPLACEMENTS = [
    # When "robots industry" appears it should be "robotics industry"
    (r'\brobots industry\b', 'robotics industry'),
    (r'\brobots ecosystem\b', 'robotics ecosystem'),
    (r'\brobots market\b', 'robotics market'),
    (r'\brobots sector\b', 'robotics sector'),
    (r'\brobots exports\b', 'robotics exports'),
    (r'\brobots export\b', 'robotics export'),
    (r'\bChinese robots\b', 'Chinese robotics'),  # usually talking about the industry
    # Category drift
    (r'\bhumanoid robots category\b', 'humanoid robotics category'),
]

def main(docx_path):
    input_path = Path(docx_path)
    output_path = input_path.with_stem(input_path.stem + "_cleaned")

    doc = Document(str(input_path))
    replace_counts = {}

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        original = para.text
        current = original

        # Apply direct replacements
        for pattern, replacement in DIRECT_REPLACEMENTS:
            new, count = re.subn(pattern, replacement, current)
            if count:
                key = f'{pattern} → {replacement}'
                replace_counts[key] = replace_counts.get(key, 0) + count
                current = new

        # Apply context replacements
        for pattern, replacement in CONTEXT_REPLACEMENTS:
            new, count = re.subn(pattern, replacement, current, flags=re.IGNORECASE)
            if count:
                key = f'{pattern} → {replacement}'
                replace_counts[key] = replace_counts.get(key, 0) + count
                current = new

        if current != original:
            # Write back preserving first-run formatting
            if para.runs:
                para.runs[0].text = current
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.text = current

    # Same for text inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    original = para.text
                    current = original
                    for pattern, replacement in DIRECT_REPLACEMENTS:
                        new, count = re.subn(pattern, replacement, current)
                        if count:
                            key = f'(table) {pattern} → {replacement}'
                            replace_counts[key] = replace_counts.get(key, 0) + count
                            current = new
                    for pattern, replacement in CONTEXT_REPLACEMENTS:
                        new, count = re.subn(pattern, replacement, current, flags=re.IGNORECASE)
                        if count:
                            key = f'(table) {pattern} → {replacement}'
                            replace_counts[key] = replace_counts.get(key, 0) + count
                            current = new
                    if current != original:
                        if para.runs:
                            para.runs[0].text = current
                            for r in para.runs[1:]:
                                r.text = ""
                        else:
                            para.text = current

    doc.save(str(output_path))

    # Report
    print(f"Saved: {output_path}")
    print()
    print("Replacements made:")
    for key, count in sorted(replace_counts.items(), key=lambda x: -x[1]):
        print(f"  {count:4d} ×  {key}")
    print()
    print(f"Total replacements: {sum(replace_counts.values())}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 cleanup_humanized.py <humanized.docx>")
        sys.exit(1)
    main(sys.argv[1])
