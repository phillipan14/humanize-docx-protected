"""Humanizer wrapper with proper-noun / citation / reference protection.

Processes a docx paragraph by paragraph through EN -> zh-CN -> ko -> EN roundtrip,
BUT protects proper nouns, citations, years, and reference-section entries so they
come back unchanged.

Paragraphs skipped entirely:
  - Empty / whitespace-only
  - Headings (based on Word style containing "Heading")
  - Any paragraph inside the "References" / "Bibliography" / "Appendix D" sections
  - Paragraphs that are primarily URL / DOI / file-path
  - Very short paragraphs (< 20 words)
  - Paragraphs dominated by citations (>30% of tokens are citations)

Within each translatable paragraph:
  1. Replace proper nouns with placeholders (§1§, §2§, ...) that translators won't paraphrase
  2. Replace (Author, Year) citations with placeholders
  3. Replace URLs / DOIs with placeholders
  4. Roundtrip translate EN -> zh-CN -> ko -> EN
  5. Restore all placeholders

Progress is saved every 5 paragraphs; use --resume to continue after interruption.

Usage:
    python3 humanize_protected.py "Capstone Draft v3.docx"
    python3 humanize_protected.py "Capstone Draft v3.docx" --resume
"""
import sys, time, random, re
from pathlib import Path
from deep_translator import GoogleTranslator
from docx import Document

MAX_CHARS = 4500
SAVE_EVERY = 5
DELAY_MIN = 1.8
DELAY_MAX = 3.2

CHAIN = [("en", "zh-CN"), ("zh-CN", "ko"), ("ko", "en")]

# Proper nouns / names / framework terms / company names we MUST protect
PROTECTED_TERMS = [
    # Team & advisor
    "Phillip Guangning An", "Phillip An", "Phillip",
    "Angelo Mok", "Angelo",
    "Boran Cui", "Boran",
    "Professor David Q. Pan", "Professor Pan", "David Pan", "David Q. Pan",
    "潘庆中", "Pan Qingzhong",
    "Jack",
    # Schools & programs
    "Schwarzman College", "Schwarzman Scholar", "Schwarzman Scholars Program", "Schwarzman Scholars", "Schwarzman",
    "Tsinghua University", "Tsinghua",
    "Master of Global Affairs", "MGA",
    # Framework & key terminology
    "Hardware–Software–Data/Service", "Hardware-Software-Data/Service", "H-S-D framework", "H-S-D",
    # Partner org & group capstone terms
    "Skylarq AI", "Skylarq",
    "Capstone Project", "Capstone", "Individual Reflection",
    # Chinese company names
    "Geekplus", "Hai Robotics", "Quicktron", "Libiao", "Mushiny",
    "Pudu Robotics", "Pudu", "KEENON", "Keenon", "Gaussian Robotics", "Gaussian", "Reeman", "Orion Star",
    "AUBO Robotics", "AUBO", "JAKA Robotics", "JAKA", "Elite Robots", "Elite", "Dobot", "Flexiv",
    "Estun Automation", "Estun", "SIASUN", "Efort", "Inovance", "GSK CNC", "ROKAE",
    "MicroPort MedBot", "MicroPort", "Tinavi", "Edge Medical", "Wego",
    "Fourier Intelligence", "Fourier",
    "Unitree Robotics", "Unitree", "Galaxy Robotics", "Galaxy", "AGI Bot", "Agibot",
    "LimX Dynamics", "LimX", "Kepler", "Galbot", "Deep Robotics", "CloudMinds",
    "IO-AI Tech", "IO-AI",
    "BitRobot Network", "BitRobot", "FrodoBox Labs", "FrodoBox",
    "Dexta Robotics", "Dexta", "Manus", "Rococo", "Mantis",
    # Foreign benchmarks
    "Universal Robots", "FANUC", "ABB", "KUKA",
    "Looking Glass XR", "Physical Intelligence", "Skild",
    "Agility Robotics", "Agility", "Figure AI", "Boston Dynamics", "SenseGlove",
    # Investors / VCs
    "Hill House Capital", "Hillhouse Capital", "Hillhouse",
    "Andreessen Horowitz", "Andreessen", "Porsche", "Lenovo", "Volkswagen", "SK Hynix", "IDG",
    # Mercor + Simovian
    "Mercor", "Simovian Intelligence", "Simovian",
    "Boundless VC", "Boundless",
    # Government / Institutional
    "International Federation of Robotics", "IFR",
    "National Bureau of Statistics of China", "NBS", "SCIO", "State Council Information Office",
    "USTR", "Bureau of Industry and Security", "BIS", "Federal Reserve Board",
    "U.S. Bureau of Labor Statistics", "BLS", "Eurostat",
    "JETRO", "KOTRA", "ITIF", "CSIS", "CSET",
    "HKEX", "Hong Kong Exchanges and Clearing",
    "FDA", "CE MDR", "CE", "ISO/TS 15066",
    "European Commission", "European Union", "Office of Foreign Assets Control", "OFAC",
    "Tsinghua University Graduate Thesis Writing Guide",
    # Tech / model
    "GPT-4", "Claude", "DeepSeek", "ChatGPT", "Gemini",
    "VLA", "UMI", "Universal Manipulation Interface",
    "Teleoperation", "AMR", "SLAM",
    # Companies mentioned in text
    "Ceres", "Neuromecca", "Dexmed", "ByteDance", "Google DeepMind", "DeepMind", "Google", "OpenAI", "Anthropic",
    "Bosch", "Imperial College London", "Momenta",
    "Tencent", "Alibaba", "BMW", "Mercedes", "Xiaomi", "GE", "Nvidia", "NVIDIA",
    # Schwarzman capstone precedents
    "Xu Jia", "Rui Kaili", "Zou Yujia", "Wang Ziqi",
    # LaTeX / Word refs
    "Thuthesis", "Tsinghua Guide to Thesis Writing",
    # Capstone artifacts
    "Executive Summary", "Table of Contents", "Abstract", "Acknowledgements",
    "Appendix A", "Appendix B", "Appendix C", "Appendix D", "Appendix E",
    "Chapter 1", "Chapter 2", "Chapter 3", "Chapter 4", "Chapter 5", "Chapter 6", "Chapter 7", "Chapter 8",
    "Section 4.3", "Section 4.4", "Section 5.14", "Section 6.3.4", "Section 6.5.1",
    # Places — only multi-word or non-country-adjective forms (single country
    # names like "China" roundtrip reliably and don't need protection).
    "Beijing", "Shanghai", "Shenzhen", "Suzhou", "Hangzhou", "Guangzhou", "Huzhou",
    "United States", "U.S.",
    "South Korea", "Singapore",
    "United Kingdom",
    "Southeast Asia", "SEA",
    "European Union",
    "Macau",
    # Book anchors often appearing with Chinese
    "摘要", "声明",
    # Additional quoted citation-like names
    "SCIO", "NBS", "IFR", "HKEX", "USTR", "OFAC", "FDA", "BIS", "EU", "FTC",
]

# Sort by length descending so longer matches get replaced first
PROTECTED_TERMS = sorted(set(PROTECTED_TERMS), key=lambda x: -len(x))

# Citation pattern: (Author, 2025) or (Author et al., 2025) or (State Council Information Office, 2026)
CITATION_PATTERN = re.compile(r'\([A-Z][A-Za-z0-9 .,&\'\-]+?,\s*(?:19|20)\d{2}[a-z]?\)')
# Interview citation: (author interview, April 5, 2026)
INTERVIEW_PATTERN = re.compile(r'\([A-Za-z ,]+interview[^)]*?202\d\)', re.IGNORECASE)
# URL
URL_PATTERN = re.compile(r'https?://\S+')
# DOI / numeric-only specific patterns
NUMERIC_BRACKET = re.compile(r'\$[\d,]+(?:–[\d,]+)?(?:/hr)?')
# Year
YEAR_PATTERN = re.compile(r'\b20\d{2}\b|\b19\d{2}\b')

# Sections to SKIP entirely (detected by heading text or context)
SKIP_HEADINGS = [
    "References", "Bibliography", "Appendix A", "Appendix B", "Appendix C", "Appendix D", "Appendix E",
    "LIST OF FIGURES", "LIST OF SYMBOLS", "TABLE OF CONTENTS", "Generative AI Use Disclosure",
    "声明", "RESUME", "Interview D.", "COMMENTS FROM", "RESOLUTION OF",
    "摘要", "Chinese Abstract",
    "LIST OF FIGURES AND TABLES", "LIST OF SYMBOLS AND ACRONYMS",
    "About Schwarzman College Group Capstone",
    "Responsibility Description",
]

def is_heading(para):
    """True if paragraph has a Heading style."""
    if para.style and para.style.name:
        return 'Heading' in para.style.name or 'Title' in para.style.name
    return False

def is_skip_paragraph(para, current_section):
    """Return True if paragraph should be skipped (no translation)."""
    text = para.text.strip()
    if not text:
        return True
    if is_heading(para):
        return True
    if current_section and any(h.lower() in current_section.lower() for h in SKIP_HEADINGS):
        return True
    if len(text.split()) < 20:
        return True
    if URL_PATTERN.search(text) and len(URL_PATTERN.findall(text)[0]) > len(text) * 0.3:
        return True
    if '[To be' in text or '[resume placeholder' in text or '[Title of' in text:
        return True
    if (text.startswith('Table ') or text.startswith('Figure ')) and len(text.split()) < 30:
        return True
    if text.isupper() and len(text.split()) < 10:
        return True
    # Paragraphs containing figure/table insertion markers — don't translate placeholder text
    if text.startswith('[FIGURE ') or text.startswith('[TABLE '):
        return True
    # Skip paragraphs where citations dominate (>3 citations or citation coverage >20% of chars)
    cit_count = len(CITATION_PATTERN.findall(text)) + len(INTERVIEW_PATTERN.findall(text))
    if cit_count > 3:
        return True
    cit_chars = sum(len(m) for m in CITATION_PATTERN.findall(text)) + sum(len(m) for m in INTERVIEW_PATTERN.findall(text))
    if cit_chars > len(text) * 0.2:
        return True
    # Skip if paragraph has multiple dollar-amount numeric sequences (pricing tables)
    if len(NUMERIC_BRACKET.findall(text)) > 2:
        return True
    return False

def protect(text):
    """Replace protected terms, citations, URLs with placeholders. Returns (text, mapping).

    Uses TK followed by zero-padded digits as placeholder format. Google Translate
    preserves short uppercase sequences well (e.g., "TK0042") better than § chars.
    Protects in order: URL, citations, proper nouns (word-bounded), years/dollars.
    """
    mapping = {}
    counter = [0]
    def make_token():
        counter[0] += 1
        return f"ZKTK{counter[0]:04d}ZK"

    # 1. Protect URLs first
    def url_sub(m):
        t = make_token()
        mapping[t] = m.group(0)
        return t
    text = URL_PATTERN.sub(url_sub, text)

    # 2. Protect APA + interview citations (regex captures them whole)
    def cit_sub(m):
        t = make_token()
        mapping[t] = m.group(0)
        return t
    text = CITATION_PATTERN.sub(cit_sub, text)
    text = INTERVIEW_PATTERN.sub(cit_sub, text)

    # 3. Protect proper nouns with WORD-BOUNDARY matching so "China" doesn't
    #    match inside "Chinese" and single-word terms don't over-replace.
    #    One token per unique term (not per match), so repeat mentions share.
    term_tokens = {}
    for term in PROTECTED_TERMS:
        if not term:
            continue
        escaped = re.escape(term)
        if term[0].isalnum() and term[-1].isalnum():
            pattern = re.compile(r'\b' + escaped + r'\b')
        else:
            pattern = re.compile(escaped)
        if pattern.search(text):
            if term not in term_tokens:
                t = make_token()
                term_tokens[term] = t
                mapping[t] = term
            else:
                t = term_tokens[term]
            text = pattern.sub(t, text)

    # 4. Dollar amounts first (they contain digits that would also match YEAR)
    text = NUMERIC_BRACKET.sub(lambda m: (lambda t: (mapping.update({t: m.group(0)}), t)[1])(make_token()), text)
    # Years last (remaining standalone years)
    text = YEAR_PATTERN.sub(lambda m: (lambda t: (mapping.update({t: m.group(0)}), t)[1])(make_token()), text)

    return text, mapping

def restore(text, mapping):
    """Restore placeholders back to protected terms.

    Because Google Translate may duplicate or slightly alter tokens, we:
    1. Collapse any near-duplicates of a token (e.g., TK0001TK0001 -> TK0001)
    2. Replace each token with its original
    3. Clean up any orphaned token-like patterns that weren't in the mapping
    """
    # Step 1: Collapse consecutive duplicates of same token
    for token in mapping:
        # Two or more times in a row -> one
        text = re.sub(r'(?:' + re.escape(token) + r')\s*(?:' + re.escape(token) + r')+', token, text)

    # Step 2: Restore — single pass, longest tokens first so nested wouldn't happen
    for token in sorted(mapping, key=lambda k: -len(k)):
        if token in text:
            text = text.replace(token, mapping[token])

    # Step 3: Clean up orphan placeholder patterns that survived (broken tokens)
    text = re.sub(r'ZKTK\d*ZK?', '', text)
    text = re.sub(r'ZK\d*\s*ZK', '', text)
    # Collapse any double-spaces introduced
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()

def translate_chunk(text, src, tgt, retries=3):
    if not text.strip():
        return text
    for attempt in range(retries):
        try:
            return GoogleTranslator(source=src, target=tgt).translate(text)
        except Exception as e:
            if attempt < retries - 1:
                wait = (attempt + 1) * 10 + random.uniform(0, 5)
                print(f"    rate-limit / error — waiting {wait:.0f}s...", file=sys.stderr)
                time.sleep(wait)
            else:
                raise

def translate_text(text, src, tgt):
    if len(text) <= MAX_CHARS:
        return translate_chunk(text, src, tgt)
    # chunk on sentence boundaries
    sents = text.replace('. ', '.\n').split('\n')
    chunks, cur = [], ""
    for s in sents:
        if len(cur) + len(s) + 1 > MAX_CHARS:
            if cur: chunks.append(cur)
            cur = s
        else:
            cur = (cur + ' ' + s).strip() if cur else s
    if cur: chunks.append(cur)
    out = []
    for c in chunks:
        out.append(translate_chunk(c, src, tgt))
        time.sleep(random.uniform(0.5, 1.2))
    return ' '.join(out)

def roundtrip(text):
    cur = text
    for src, tgt in CHAIN:
        cur = translate_text(cur, src, tgt)
        time.sleep(random.uniform(0.8, 1.5))
    return cur

def main(docx_path, resume=False):
    input_path = Path(docx_path)
    output_path = input_path.with_stem(input_path.stem + "_humanized")

    if resume and output_path.exists():
        doc = Document(str(output_path))
        original_doc = Document(str(input_path))
        done = set()
        for i, (o, n) in enumerate(zip(original_doc.paragraphs, doc.paragraphs)):
            if o.text.strip() and o.text != n.text:
                done.add(i)
        print(f"Resuming: {len(done)} paragraphs already humanized.\n")
    else:
        doc = Document(str(input_path))
        done = set()

    # Track current section heading as we walk paragraphs
    current_section = ""
    all_paras = list(enumerate(doc.paragraphs))
    total = len(all_paras)
    processed = 0

    for i, para in all_paras:
        # Update section tracker
        if is_heading(para) and para.text.strip():
            current_section = para.text.strip()
            continue

        if i in done:
            continue

        if is_skip_paragraph(para, current_section):
            continue

        text = para.text.strip()
        label = text[:80].replace('\n', ' ')
        print(f"[§ {current_section[:40]} | para {i}/{total}] {label}...")

        try:
            protected, mapping = protect(text)
            translated = roundtrip(protected)
            restored = restore(translated, mapping)
        except Exception as e:
            print(f"  ERROR: {e} — keeping original")
            doc.save(str(output_path))
            continue

        # Replace paragraph text while preserving formatting of first run
        if para.runs:
            para.runs[0].text = restored
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.text = restored

        processed += 1
        print(f"  done ({len(text)} → {len(restored)} chars, {len(mapping)} protections)")

        if processed % SAVE_EVERY == 0:
            doc.save(str(output_path))
            print(f"  [progress: {processed} paragraphs humanized, saved]")

        time.sleep(random.uniform(DELAY_MIN, DELAY_MAX))

    doc.save(str(output_path))
    print(f"\nCompleted. Output: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 humanize_protected.py <document.docx> [--resume]")
        sys.exit(1)
    resume = "--resume" in sys.argv
    main(sys.argv[1], resume=resume)
